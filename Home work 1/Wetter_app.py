from docx import Document

# Создание нового документа
doc = Document()

# Немецкая версия
doc.add_heading('Technische Spezifikation: Wetter-Tracker-App', level=1)

doc.add_heading('Zusammenfassung', level=2)
doc.add_paragraph(
    'Die Wetter-Tracker-App zeigt den aktuellen Wetterzustand und eine 7-Tage-Prognose für beliebige Orte an. '
    'Nutzer können bevorzugte Städte speichern und erhalten auf Wunsch Benachrichtigungen bei extremen Wetterlagen. '
    'So behalten sie stets den Überblick über das Wetter an wichtigen Orten.'
)

doc.add_heading('Hintergrund', level=2)
doc.add_paragraph(
    'Wetter-Apps sind oft überladen oder unübersichtlich. Viele Nutzer wünschen sich eine einfache, schnelle Möglichkeit, '
    'aktuelle Wetterdaten sowie Vorhersagen für ausgewählte Orte abzurufen, ohne von Werbung oder komplexen Funktionen abgelenkt zu werden. '
    'Die Wetter-Tracker-App löst dieses Problem durch eine minimalistische, benutzerfreundliche Gestaltung.'
)

doc.add_heading('Ziele', level=2)
doc.add_paragraph('1. Aktuelle Wetterdaten und eine 7-Tage-Vorhersage für gespeicherte Orte anzeigen.')
doc.add_paragraph('2. Benutzer können ihre Lieblingsstädte hinzufügen und verwalten.')
doc.add_paragraph('3. Push-Benachrichtigungen bei extremen Wetterwarnungen (optional einstellbar).')

doc.add_heading('Nicht-Ziele', level=2)
doc.add_paragraph('Die App wird keine eigenen Wetterdaten generieren oder meteorologische Analysen erstellen.')
doc.add_paragraph('Es wird keine soziale Netzwerkfunktion (z.B. Teilen von Wetterinformationen) geben.')

doc.add_heading('Stakeholder', level=2)
doc.add_paragraph('Reisende und Berufspendler: Nutzer, die das Wetter an mehreren Orten im Blick behalten möchten.')
doc.add_paragraph('Outdoor-Enthusiasten: Menschen, die Aktivitäten im Freien planen und auf präzise Wetterinformationen angewiesen sind.')

doc.add_heading('Technische Beschreibung (optional)', level=2)
doc.add_paragraph(
    '- Frontend: Native App-Entwicklung (Flutter oder React Native) für iOS und Android.\n'
    '- Backend: Nutzung einer externen Wetter-API (z.B. OpenWeatherMap oder WeatherAPI).\n'
    '- Benachrichtigungen: Integration von Push-Services wie Firebase Cloud Messaging.\n'
    '- Datenspeicherung: Lokale Speicherung der Favoritenorte im Gerät oder in der Cloud (Google Firebase).'
)

# Переход на русскую версию
doc.add_page_break()

doc.add_heading('Техническая спецификация: Приложение "Погодный трекер"', level=1)

doc.add_heading('Краткое описание', level=2)
doc.add_paragraph(
    'Приложение "Погодный трекер" показывает текущие погодные условия и прогноз на 7 дней для выбранных пользователем мест. '
    'Пользователи могут сохранять избранные города и получать уведомления о чрезвычайных погодных ситуациях. '
    'Это помогает всегда быть в курсе погодных изменений в важных локациях.'
)

doc.add_heading('Предпосылки', level=2)
doc.add_paragraph(
    'Многие погодные приложения перегружены рекламой и лишними функциями. Пользователи хотят иметь простое и быстрое средство '
    'для получения точной информации о погоде, без отвлекающих элементов. Приложение "Погодный трекер" решает эту проблему '
    'благодаря минималистичному и удобному дизайну.'
)

doc.add_heading('Цели', level=2)
doc.add_paragraph('1. Отображение текущей погоды и 7-дневного прогноза для сохранённых мест.')
doc.add_paragraph('2. Возможность добавлять и управлять списком любимых городов.')
doc.add_paragraph('3. Отправка push-уведомлений о предупреждениях при экстремальных погодных условиях (по желанию пользователя).')

doc.add_heading('Не-цели', level=2)
doc.add_paragraph('Приложение не будет самостоятельно генерировать погодные данные или проводить метеорологические анализы.')
doc.add_paragraph('Приложение не будет иметь социальных функций, например, публикации или обмена погодной информацией.')

doc.add_heading('Заинтересованные стороны', level=2)
doc.add_paragraph('Путешественники и ежедневные пассажиры: пользователи, которым нужно следить за погодой в разных местах.')
doc.add_paragraph('Любители активного отдыха: люди, планирующие мероприятия на открытом воздухе и нуждающиеся в точном прогнозе погоды.')

doc.add_heading('Техническое описание (опционально)', level=2)
doc.add_paragraph(
    '- Фронтенд: нативная разработка приложения с использованием Flutter или React Native для iOS и Android.\n'
    '- Бэкенд: использование внешнего погодного API (например, OpenWeatherMap или WeatherAPI).\n'
    '- Уведомления: интеграция службы push-уведомлений, например, через Firebase Cloud Messaging.\n'
    '- Хранение данных: локальное хранение списка любимых городов на устройстве пользователя или в облаке (например, Google Firebase).'
)

# Сохраняем файл
file_path = '/mnt/data/Wetter_Tracker_Spec.docx'
doc.save(file_path)

file_path
